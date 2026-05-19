from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Any
import xml.etree.ElementTree as ET
import zipfile

from app.core.canonical_artifacts import (
    ArtifactLineageRef,
    CanonicalArtifactBundle,
    CanonicalLayoutBlock,
    CanonicalSourceManifest,
    CanonicalTabularFrame,
    CanonicalTextBlock,
)
from app.core.config import settings
from app.core.structured_logging import emit_structured_log
from app.services.canonical_frame_consolidator import consolidate_fragmented_frames
from app.services.canonical_native_tabular_ingestion import (
    build_native_frame_payload,
    extract_native_tabular_frames,
)
from app.services.file_intelligence_router import FileIntelligenceRouter
from app.services.ocr_layout_gateway import extract_image_layout_payload


def _decode_text_bytes(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return ""


def _append_manifest_warning(manifest: CanonicalSourceManifest, message: str) -> None:
    if message and message not in manifest.warnings:
        manifest.warnings.append(message)


def _build_synthetic_column_names(column_count: int) -> list[str]:
    return [f"column_{index}" for index in range(1, max(column_count, 0) + 1)]


def _normalize_table_rows(raw_rows: list[list[Any]]) -> list[list[str]]:
    normalized_rows: list[list[str]] = []
    for row in raw_rows:
        normalized_row = [str(cell or "").strip() for cell in list(row or [])]
        if any(normalized_row):
            normalized_rows.append(normalized_row)
    return normalized_rows


def _pick_table_column_names(rows: list[list[str]]) -> list[str]:
    if not rows:
        return []
    header = rows[0]
    if not any(header):
        return _build_synthetic_column_names(max(len(row) for row in rows))
    return [
        value if value else f"column_{index}"
        for index, value in enumerate(header, start=1)
    ]


def _docx_text_from_element(element: ET.Element, namespace: dict[str, str]) -> str:
    text_parts = []
    for node in element.findall(".//w:t", namespace):
        if node.text:
            text_parts.append(str(node.text))
    return "".join(text_parts).strip()


def _extract_docx_openxml_payload(file_bytes: bytes) -> tuple[list[str], list[list[list[str]]]]:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
        document_xml = archive.read("word/document.xml")

    root = ET.fromstring(document_xml)
    body = root.find("w:body", namespace)
    if body is None:
        return [], []

    paragraphs: list[str] = []
    tables: list[list[list[str]]] = []
    for child in list(body):
        tag_name = str(child.tag)
        if tag_name.endswith("}p"):
            text = _docx_text_from_element(child, namespace)
            if text:
                paragraphs.append(text)
            continue
        if tag_name.endswith("}tbl"):
            table_rows: list[list[str]] = []
            for row in child.findall("./w:tr", namespace):
                cells: list[str] = []
                for cell in row.findall("./w:tc", namespace):
                    cells.append(_docx_text_from_element(cell, namespace))
                if any(str(value or "").strip() for value in cells):
                    table_rows.append(cells)
            if table_rows:
                tables.append(table_rows)
    return paragraphs, tables


def _extract_pdf_text_tables_from_blocks(bundle: CanonicalArtifactBundle) -> None:
    if bundle.tabular_frames:
        return

    frame_index = 0
    for block in bundle.text_blocks:
        lines = [str(line or "").strip() for line in str(block.text or "").splitlines()]
        candidate_groups: list[list[list[str]]] = []
        current_group: list[list[str]] = []
        for line in lines:
            if "|" in line:
                columns = [segment.strip() for segment in line.split("|")]
                if sum(1 for value in columns if value) >= 2:
                    current_group.append(columns)
                    continue
            if len(current_group) >= 2:
                candidate_groups.append(current_group)
            current_group = []
        if len(current_group) >= 2:
            candidate_groups.append(current_group)

        for group in candidate_groups:
            normalized_rows = _normalize_table_rows(group)
            if len(normalized_rows) < 2:
                continue
            frame_index += 1
            column_names = _pick_table_column_names(normalized_rows)
            body_rows = normalized_rows[1:] if len(normalized_rows) > 1 else normalized_rows
            bundle.tabular_frames.append(
                CanonicalTabularFrame(
                    frame_id=f"pdf-text-table-{frame_index}",
                    label=f"PDF Text Table #{frame_index}",
                    row_count=len(body_rows),
                    column_count=max((len(row) for row in normalized_rows), default=0),
                    column_names=column_names or _build_synthetic_column_names(max((len(row) for row in normalized_rows), default=0)),
                    extraction_confidence=0.55,
                    lineage=[
                        ArtifactLineageRef(
                            page_number=block.page_number,
                            table_id=f"pdf-text-table-{frame_index}",
                            row_start=1 if body_rows else None,
                            row_end=len(body_rows) if body_rows else None,
                            column_names=column_names,
                            metadata={"source_kind": "pdf_text_heuristic"},
                        )
                    ],
                    metadata={
                        "source_kind": "pdf_text_table",
                        "header_detected": bool(normalized_rows and any(normalized_rows[0])),
                        "sample_rows": body_rows[:3],
                    },
                )
            )


@dataclass
class ArtifactParserResult:
    bundle: CanonicalArtifactBundle
    parser_name: str


class BaseArtifactParser:
    parser_name = "base"

    @classmethod
    def parse(
        cls,
        *,
        manifest: CanonicalSourceManifest,
        file_bytes: bytes,
        metadata: dict[str, Any] | None = None,
    ) -> ArtifactParserResult:
        bundle = CanonicalArtifactBundle(
            source_manifest=manifest,
            extraction_confidence=manifest.detection.detection_confidence,
            metadata=metadata or {},
        )
        return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)


class DelegatedTabularArtifactParser(BaseArtifactParser):
    parser_name = "delegated_tabular"

    @classmethod
    def parse(
        cls,
        *,
        manifest: CanonicalSourceManifest,
        file_bytes: bytes,
        metadata: dict[str, Any] | None = None,
    ) -> ArtifactParserResult:
        bundle = CanonicalArtifactBundle(
            source_manifest=manifest,
            tabular_frames=[
                CanonicalTabularFrame(
                    frame_id="legacy-tabular-runtime",
                    label="Delegated to legacy DataEngine",
                    extraction_confidence=1.0,
                    metadata={
                        "delegated": True,
                        "runtime_owner": "legacy_data_engine",
                    },
                )
            ],
            extraction_confidence=1.0,
            metadata=metadata or {},
        )
        _append_manifest_warning(
            bundle.source_manifest,
            "La extracción tabular sigue delegada al runtime activo hasta completar la integración enterprise.",
        )
        return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)


class NativeTabularArtifactParser(BaseArtifactParser):
    parser_name = "native_tabular_parallel"

    @classmethod
    def parse(
        cls,
        *,
        manifest: CanonicalSourceManifest,
        file_bytes: bytes,
        metadata: dict[str, Any] | None = None,
    ) -> ArtifactParserResult:
        bundle = CanonicalArtifactBundle(
            source_manifest=manifest,
            extraction_confidence=0.0,
            metadata=metadata or {},
        )

        try:
            extracted_frames = extract_native_tabular_frames(
                file_name=manifest.file_name,
                file_bytes=file_bytes,
                extension=str(manifest.extension or "").lower(),
            )
        except Exception as exc:
            _append_manifest_warning(
                bundle.source_manifest,
                f"La extracción tabular nativa paralela falló de forma segura: {exc}",
            )
            return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)

        for frame in extracted_frames:
            frame_payload = build_native_frame_payload(frame)
            frame_id = frame_payload["frame_id"]
            sheet_name = frame_payload["metadata"].get("sheet_name")
            if frame_payload["metadata"].get("truncated_rows"):
                _append_manifest_warning(
                    bundle.source_manifest,
                    f"La extracción tabular paralela truncó filas en '{sheet_name or frame_payload['label']}' para mantener límites de seguridad.",
                )
            if frame_payload["metadata"].get("truncated_columns"):
                _append_manifest_warning(
                    bundle.source_manifest,
                    f"La extracción tabular paralela truncó columnas en '{sheet_name or frame_payload['label']}' para mantener límites de seguridad.",
                )
            bundle.tabular_frames.append(
                CanonicalTabularFrame(
                    frame_id=frame_id,
                    label=frame_payload["label"],
                    row_count=frame_payload["row_count"],
                    column_count=frame_payload["column_count"],
                    column_names=frame_payload["column_names"],
                    extraction_confidence=1.0,
                    lineage=[
                        ArtifactLineageRef(
                            sheet_name=str(sheet_name or "") or None,
                            row_start=1 if frame_payload["row_count"] else None,
                            row_end=frame_payload["row_count"] if frame_payload["row_count"] else None,
                            column_names=frame_payload["column_names"],
                            metadata={"source_kind": manifest.source_kind.value},
                        )
                    ],
                    metadata={
                        **frame_payload["metadata"],
                        "source_kind": manifest.source_kind.value,
                        "delegated": False,
                    },
                )
            )

        if not bundle.tabular_frames:
            _append_manifest_warning(
                bundle.source_manifest,
                "La extracción tabular nativa no encontró hojas/filas aprovechables.",
            )
        else:
            bundle.extraction_confidence = 1.0
            bundle.metadata.update(
                {
                    "frame_count": len(bundle.tabular_frames),
                    "parser_name": cls.parser_name,
                }
            )
        return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)


class PlainTextArtifactParser(BaseArtifactParser):
    parser_name = "plain_text"

    @classmethod
    def parse(
        cls,
        *,
        manifest: CanonicalSourceManifest,
        file_bytes: bytes,
        metadata: dict[str, Any] | None = None,
    ) -> ArtifactParserResult:
        text = _decode_text_bytes(file_bytes).strip()
        bundle = CanonicalArtifactBundle(
            source_manifest=manifest,
            extraction_confidence=1.0 if text else 0.0,
            metadata=metadata or {},
        )
        if text:
            bundle.text_blocks.append(
                CanonicalTextBlock(
                    block_id="text-block-1",
                    text=text,
                    block_type="plain_text",
                    extraction_confidence=1.0,
                    lineage=[ArtifactLineageRef(column_names=[])],
                )
            )
        else:
            _append_manifest_warning(bundle.source_manifest, "No se pudo decodificar texto legible del archivo plano.")
        return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)


class PdfArtifactParser(BaseArtifactParser):
    parser_name = "pdf_text"

    @classmethod
    def _extract_pdf_tables(
        cls,
        *,
        bundle: CanonicalArtifactBundle,
        file_bytes: bytes,
    ) -> None:
        if not settings.FILE_INTELLIGENCE_ENABLE_PDF_TABLE_EXTRACTION:
            return

        try:
            import pdfplumber
        except ImportError:
            _append_manifest_warning(
                bundle.source_manifest,
                "La extracción tabular de PDF está deshabilitada por dependencia faltante ('pdfplumber').",
            )
            return

        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf_document:
                for page_index, page in enumerate(pdf_document.pages, start=1):
                    tables = page.extract_tables() or []
                    for table_index, raw_table in enumerate(tables, start=1):
                        rows = _normalize_table_rows(raw_table or [])
                        if not rows:
                            continue
                        column_names = _pick_table_column_names(rows)
                        body_rows = rows[1:] if len(rows) > 1 else rows
                        row_count = len(body_rows)
                        column_count = max((len(row) for row in rows), default=0)
                        bundle.tabular_frames.append(
                            CanonicalTabularFrame(
                                frame_id=f"pdf-table-{page_index}-{table_index}",
                                label=f"PDF Table Page {page_index} #{table_index}",
                                row_count=row_count,
                                column_count=column_count,
                                column_names=column_names or _build_synthetic_column_names(column_count),
                                extraction_confidence=0.78,
                                lineage=[
                                    ArtifactLineageRef(
                                        page_number=page_index,
                                        table_id=f"pdf-table-{page_index}-{table_index}",
                                        row_start=1 if row_count else None,
                                        row_end=row_count if row_count else None,
                                        column_names=column_names,
                                        metadata={"source_kind": "pdf"},
                                    )
                                ],
                                metadata={
                                    "source_kind": "pdf_table",
                                    "header_detected": bool(rows and any(rows[0])),
                                    "sample_rows": body_rows[:3],
                                },
                            )
                        )
                bundle.tabular_frames = consolidate_fragmented_frames(bundle.tabular_frames)
        except Exception as exc:
            _append_manifest_warning(
                bundle.source_manifest,
                f"La extracción tabular del PDF falló de forma segura: {exc}",
            )

    @classmethod
    def parse(
        cls,
        *,
        manifest: CanonicalSourceManifest,
        file_bytes: bytes,
        metadata: dict[str, Any] | None = None,
    ) -> ArtifactParserResult:
        bundle = CanonicalArtifactBundle(
            source_manifest=manifest,
            extraction_confidence=0.0,
            metadata=metadata or {},
        )
        try:
            from pypdf import PdfReader
        except ImportError:
            _append_manifest_warning(bundle.source_manifest, "Falta 'pypdf' para extracción enterprise de PDF.")
            return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)

        try:
            reader = PdfReader(io.BytesIO(file_bytes))
        except Exception as exc:
            _append_manifest_warning(bundle.source_manifest, f"No se pudo abrir el PDF: {exc}")
            return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)

        extracted_pages = 0
        total_pages = len(reader.pages)
        for page_index, page in enumerate(reader.pages, start=1):
            try:
                raw_text = (page.extract_text() or "").strip()
            except Exception:
                raw_text = ""

            if raw_text:
                extracted_pages += 1
                bundle.text_blocks.append(
                    CanonicalTextBlock(
                        block_id=f"pdf-page-{page_index}",
                        text=raw_text,
                        block_type="page_text",
                        page_number=page_index,
                        extraction_confidence=1.0,
                        lineage=[
                            ArtifactLineageRef(
                                page_number=page_index,
                                metadata={"source_kind": "pdf"},
                            )
                        ],
                    )
                )
            else:
                bundle.layout_blocks.append(
                    CanonicalLayoutBlock(
                        block_id=f"pdf-page-layout-{page_index}",
                        block_type="page_without_text",
                        page_number=page_index,
                        extraction_confidence=0.0,
                        metadata={"text_extractable": False},
                    )
                )

        cls._extract_pdf_tables(bundle=bundle, file_bytes=file_bytes)
        _extract_pdf_text_tables_from_blocks(bundle)
        bundle.extraction_confidence = (extracted_pages / total_pages) if total_pages else 0.0
        bundle.metadata.update(
            {
                "page_count": total_pages,
                "pages_extracted": extracted_pages,
                "table_count": len(bundle.tabular_frames),
                "parser_name": cls.parser_name,
            }
        )
        if not bundle.text_blocks:
            _append_manifest_warning(
                bundle.source_manifest,
                "El PDF no expuso texto legible; necesitará OCR/layout en la siguiente fase.",
            )
        return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)


class DocxArtifactParser(BaseArtifactParser):
    parser_name = "docx_openxml"

    @classmethod
    def parse(
        cls,
        *,
        manifest: CanonicalSourceManifest,
        file_bytes: bytes,
        metadata: dict[str, Any] | None = None,
    ) -> ArtifactParserResult:
        bundle = CanonicalArtifactBundle(
            source_manifest=manifest,
            extraction_confidence=0.0,
            metadata=metadata or {},
        )
        paragraph_count = 0
        table_count = 0

        paragraphs: list[str] = []
        tables: list[list[list[str]]] = []
        parser_backend = "openxml_fallback"

        try:
            from docx import Document

            document = Document(io.BytesIO(file_bytes))
            parser_backend = "python_docx"
            paragraphs = [str(paragraph.text or "").strip() for paragraph in document.paragraphs if str(paragraph.text or "").strip()]
            for table in document.tables:
                raw_rows: list[list[str]] = []
                for row in table.rows:
                    cells = [str(cell.text or "").strip() for cell in row.cells]
                    if any(cells):
                        raw_rows.append(cells)
                if raw_rows:
                    tables.append(raw_rows)
        except Exception:
            try:
                paragraphs, tables = _extract_docx_openxml_payload(file_bytes)
            except Exception as exc:
                _append_manifest_warning(bundle.source_manifest, f"No se pudo abrir el DOCX: {exc}")
                return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)

        for paragraph_index, text in enumerate(paragraphs, start=1):
            if not text:
                continue
            paragraph_count += 1
            bundle.text_blocks.append(
                CanonicalTextBlock(
                    block_id=f"docx-paragraph-{paragraph_index}",
                    text=text,
                    block_type="paragraph",
                    extraction_confidence=1.0,
                    lineage=[
                        ArtifactLineageRef(
                            metadata={"paragraph_index": paragraph_index, "source_kind": "docx"},
                        )
                    ],
                    metadata={"source_kind": "docx", "paragraph_index": paragraph_index},
                )
            )
            bundle.layout_blocks.append(
                CanonicalLayoutBlock(
                    block_id=f"docx-layout-paragraph-{paragraph_index}",
                    block_type="paragraph",
                    text_excerpt=text[:240],
                    extraction_confidence=1.0,
                    metadata={"source_kind": "docx", "paragraph_index": paragraph_index},
                )
            )

        for table_index, raw_rows in enumerate(tables, start=1):
            if not raw_rows:
                continue
            rows = [" | ".join(cells) for cells in raw_rows]
            column_names = _pick_table_column_names(raw_rows)
            body_rows = raw_rows[1:] if len(raw_rows) > 1 else raw_rows
            column_count = max((len(row) for row in raw_rows), default=0)
            if not rows:
                continue
            table_count += 1
            table_text = "\n".join(rows)
            bundle.text_blocks.append(
                CanonicalTextBlock(
                    block_id=f"docx-table-{table_index}",
                    text=table_text,
                    block_type="table_text",
                    extraction_confidence=0.95,
                    lineage=[
                        ArtifactLineageRef(
                            table_id=f"table-{table_index}",
                            metadata={"source_kind": "docx"},
                        )
                    ],
                    metadata={"source_kind": "docx", "table_index": table_index},
                )
            )
            bundle.tabular_frames.append(
                CanonicalTabularFrame(
                    frame_id=f"docx-table-{table_index}",
                    label=f"DOCX Table #{table_index}",
                    row_count=len(body_rows),
                    column_count=column_count,
                    column_names=column_names or _build_synthetic_column_names(column_count),
                    extraction_confidence=0.95,
                    lineage=[
                        ArtifactLineageRef(
                            table_id=f"docx-table-{table_index}",
                            row_start=1 if body_rows else None,
                            row_end=len(body_rows) if body_rows else None,
                            column_names=column_names,
                            metadata={"source_kind": "docx"},
                        )
                    ],
                    metadata={
                        "source_kind": "docx_table",
                        "header_detected": bool(raw_rows and any(raw_rows[0])),
                        "sample_rows": body_rows[:3],
                    },
                )
            )
            bundle.layout_blocks.append(
                CanonicalLayoutBlock(
                    block_id=f"docx-layout-table-{table_index}",
                    block_type="table",
                    text_excerpt=table_text[:240],
                    extraction_confidence=0.95,
                    metadata={
                        "source_kind": "docx",
                        "table_index": table_index,
                        "row_count": len(body_rows),
                        "column_count": column_count,
                    },
                )
            )

        total_blocks = paragraph_count + table_count
        bundle.extraction_confidence = 1.0 if total_blocks else 0.0
        bundle.metadata.update(
            {
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "tabular_frame_count": len(bundle.tabular_frames),
                "layout_block_count": len(bundle.layout_blocks),
                "docx_parser_backend": parser_backend,
                "parser_name": cls.parser_name,
            }
        )
        if not bundle.text_blocks:
            _append_manifest_warning(bundle.source_manifest, "El DOCX no contenía bloques textuales aprovechables.")
        return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)


class LegacyDocArtifactParser(BaseArtifactParser):
    parser_name = "legacy_doc_conversion_gate"

    @classmethod
    def parse(
        cls,
        *,
        manifest: CanonicalSourceManifest,
        file_bytes: bytes,
        metadata: dict[str, Any] | None = None,
    ) -> ArtifactParserResult:
        bundle = CanonicalArtifactBundle(
            source_manifest=manifest,
            extraction_confidence=0.0,
            metadata=metadata or {},
        )
        _append_manifest_warning(
            bundle.source_manifest,
            "DOC binario no se procesa en línea; queda marcado para conversión controlada.",
        )
        return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)


class ImageArtifactParser(BaseArtifactParser):
    parser_name = "ocr_layout_gate"

    @classmethod
    def parse(
        cls,
        *,
        manifest: CanonicalSourceManifest,
        file_bytes: bytes,
        metadata: dict[str, Any] | None = None,
    ) -> ArtifactParserResult:
        bundle = CanonicalArtifactBundle(
            source_manifest=manifest,
            extraction_confidence=0.0,
            metadata=metadata or {},
        )
        ocr_payload = extract_image_layout_payload(file_bytes)
        bundle.metadata.update(
            {
                "ocr_backend": ocr_payload.get("backend"),
                "parser_name": cls.parser_name,
            }
        )
        for warning in list(ocr_payload.get("warnings") or []):
            _append_manifest_warning(bundle.source_manifest, str(warning))

        text = str(ocr_payload.get("text") or "").strip()
        if text:
            bundle.text_blocks.append(
                CanonicalTextBlock(
                    block_id="image-ocr-1",
                    text=text,
                    block_type="ocr_text",
                    extraction_confidence=float(ocr_payload.get("confidence") or 0.0),
                    lineage=[ArtifactLineageRef(metadata={"source_kind": "image"})],
                    metadata={"ocr_backend": ocr_payload.get("backend")},
                )
            )

        for block in list(ocr_payload.get("blocks") or []):
            block_text = str(block.get("text") or "").strip()
            bundle.layout_blocks.append(
                CanonicalLayoutBlock(
                    block_id=str(block.get("block_id") or f"ocr-block-{len(bundle.layout_blocks) + 1}"),
                    block_type=str(block.get("block_type") or "ocr_word"),
                    bbox=[float(value) for value in list(block.get("bbox") or [])],
                    text_excerpt=block_text[:240] if block_text else None,
                    extraction_confidence=float(block.get("confidence") or 0.0),
                    metadata={"source_kind": "image", "ocr_backend": ocr_payload.get("backend")},
                )
            )

        for table_index, table in enumerate(list(ocr_payload.get("tables") or []), start=1):
            column_names = [str(value or "").strip() for value in list(table.get("column_names") or [])]
            sample_rows = [list(row or []) for row in list(table.get("sample_rows") or [])]
            frame_id = str(table.get("table_id") or f"ocr-table-{table_index}")
            bundle.tabular_frames.append(
                CanonicalTabularFrame(
                    frame_id=frame_id,
                    label=f"OCR Table #{table_index}",
                    row_count=int(table.get("row_count") or 0),
                    column_count=int(table.get("column_count") or len(column_names)),
                    column_names=column_names or _build_synthetic_column_names(int(table.get("column_count") or 0)),
                    extraction_confidence=float(table.get("confidence") or ocr_payload.get("confidence") or 0.0),
                    lineage=[
                        ArtifactLineageRef(
                            table_id=frame_id,
                            row_start=1 if int(table.get("row_count") or 0) else None,
                            row_end=int(table.get("row_count") or 0) if int(table.get("row_count") or 0) else None,
                            column_names=column_names,
                            metadata={"source_kind": "image", "ocr_backend": ocr_payload.get("backend")},
                        )
                    ],
                    metadata={
                        "source_kind": "ocr_table",
                        "header_detected": bool(table.get("metadata", {}).get("header_detected", True)),
                        "sample_rows": sample_rows,
                        "ocr_backend": ocr_payload.get("backend"),
                    },
                )
            )

        bundle.extraction_confidence = float(ocr_payload.get("confidence") or 0.0)
        bundle.metadata["table_count"] = len(bundle.tabular_frames)
        return ArtifactParserResult(bundle=bundle, parser_name=cls.parser_name)


class ArtifactParserRegistry:
    """
    Registro enterprise de adaptadores.
    Se mantiene desacoplado del runtime activo hasta activar la integración.
    """

    @staticmethod
    def parse_to_bundle(
        *,
        file_name: str,
        file_bytes: bytes,
        mime_type: str | None = None,
        size_bytes: int | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> CanonicalArtifactBundle:
        manifest = FileIntelligenceRouter.inspect(
            file_name=file_name,
            file_bytes=file_bytes,
            mime_type=mime_type,
            size_bytes=size_bytes,
        )
        parser = ArtifactParserRegistry._resolve_parser(manifest)
        result = parser.parse(
            manifest=manifest,
            file_bytes=file_bytes,
            metadata=metadata,
        )
        result.bundle.metadata.setdefault("parser_name", result.parser_name)
        emit_structured_log(
            "artifact_parser_registry_bundle_built",
            file_name=file_name,
            parser_name=result.parser_name,
            support_level=result.bundle.source_manifest.support_level.value,
            preferred_mode=result.bundle.source_manifest.preferred_mode.value,
            text_block_count=len(result.bundle.text_blocks),
            table_frame_count=len(result.bundle.tabular_frames),
            layout_block_count=len(result.bundle.layout_blocks),
            extraction_confidence=result.bundle.extraction_confidence,
        )
        return result.bundle

    @staticmethod
    def _resolve_parser(manifest: CanonicalSourceManifest) -> type[BaseArtifactParser]:
        extension = (manifest.extension or "").lower()
        if manifest.analytics_ready:
            if settings.CANONICAL_NATIVE_TABULAR_EXTRACTION_ENABLED and extension in {"csv", "xlsx", "xls"}:
                return NativeTabularArtifactParser
            return DelegatedTabularArtifactParser
        if extension == "pdf":
            return PdfArtifactParser
        if extension == "docx":
            return DocxArtifactParser
        if extension == "doc":
            return LegacyDocArtifactParser
        if manifest.source_kind.value == "image":
            return ImageArtifactParser
        if manifest.source_kind.value == "plain_text":
            return PlainTextArtifactParser
        return BaseArtifactParser
